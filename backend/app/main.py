import json
import os
import socket
import urllib.error
import urllib.parse
import urllib.request
from contextlib import asynccontextmanager
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import PlainTextResponse, Response
from pydantic import BaseModel, ConfigDict, Field

from .database import GENERIC_TABLES, connect, init_db, new_id, row_to_dict, rows_to_dicts, utc_now
from .storage import ensure_project_dirs, project_root, require_project, safe_wiki_path


@asynccontextmanager
async def lifespan(_app: FastAPI):
    init_app()
    yield


app = FastAPI(title="AI 小说创作平台", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ProjectIn(BaseModel):
    title: str
    topic: str = ""
    genre: str = ""
    audience: str = ""
    tone: str = ""
    target_chapter_count: int = 0
    target_words_per_chapter: int = 0
    logline: str = ""
    synopsis: str = ""
    global_summary: str = ""
    privacy_mode: bool = True


class DeleteProjectIn(BaseModel):
    password: str = ""


class ChapterIn(BaseModel):
    outline_id: str = ""
    chapter_number: int = 1
    title: str = ""
    brief: str = ""
    draft: str = ""
    summary: str = ""
    status: str = "draft"


class VersionIn(BaseModel):
    label: str = ""
    content: str
    model: str = ""
    context_summary: str = ""


class WikiWriteIn(BaseModel):
    path: str
    content: str
    source_chapter_id: str = ""


class GenericIn(BaseModel):
    title: str = ""
    category: str = ""
    content: str = ""
    payload: dict[str, Any] = Field(default_factory=dict)
    status: str = "active"


class AiWorkflowIn(BaseModel):
    model_config = ConfigDict(extra="allow")

    chapter_id: str = ""
    prompt: str = ""
    content: str = ""
    count: int = 2
    payload: dict[str, Any] = Field(default_factory=dict)


class ModelConnectionTestIn(BaseModel):
    provider: str = "OpenAI"
    api_key: str = ""
    base_url: str = "https://api.openai.com/v1"
    model_name: str = ""
    temperature: float = 0.1
    max_tokens: int = 16


AUTH_SESSION: dict[str, Any] | None = None
OAUTH_PROVIDERS = {"openai", "github", "google", "custom"}


def init_app() -> None:
    init_db()


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


def auth_status_payload() -> dict[str, Any]:
    if AUTH_SESSION:
        return {
            "mode": "cloud",
            "authenticated": True,
            "user": AUTH_SESSION,
            "sync_enabled": False,
            "message": "已登录。小说项目仍默认保存在本机，未主动同步不会上传正文。",
        }
    return {
        "mode": "local",
        "authenticated": False,
        "user": None,
        "sync_enabled": False,
        "message": "本地模式：无需登录即可完整使用项目、章节、记忆、导出和本地 API 配置。",
    }


def normalize_oauth_provider(provider: str) -> str:
    normalized = provider.lower()
    if normalized not in OAUTH_PROVIDERS:
        raise HTTPException(status_code=404, detail="Unsupported OAuth provider")
    return normalized


@app.get("/api/auth/status")
def auth_status() -> dict[str, Any]:
    return auth_status_payload()


@app.get("/api/auth/oauth/{provider}/start")
def oauth_start(provider: str) -> dict[str, Any]:
    normalized = normalize_oauth_provider(provider)
    return {
        "provider": normalized,
        "requires_redirect": False,
        "authorization_url": f"/api/auth/oauth/{normalized}/callback?code=mock-code",
        "state": "local-mvp-mock-state",
        "message": "本地 MVP 暂不跳转第三方授权；后续可在这里接入真实 OAuth Provider。",
        "available_providers": sorted(OAUTH_PROVIDERS),
    }


@app.get("/api/auth/oauth/{provider}/callback")
def oauth_callback(provider: str, code: str = "", state: str = "") -> dict[str, Any]:
    global AUTH_SESSION
    normalized = normalize_oauth_provider(provider)
    AUTH_SESSION = {
        "id": f"mock-{normalized}-user",
        "provider": normalized,
        "name": "本地 OAuth 预览用户",
        "email": f"writer@{normalized}.local",
        "avatar_url": "",
    }
    return auth_status_payload() | {"oauth_code_received": bool(code), "state": state}


@app.post("/api/auth/logout")
def auth_logout() -> dict[str, Any]:
    global AUTH_SESSION
    AUTH_SESSION = None
    return auth_status_payload()


@app.post("/api/projects")
def create_project(payload: ProjectIn) -> dict[str, Any]:
    project_id = new_id()
    root = ensure_project_dirs(project_id)
    now = utc_now()
    with connect() as conn:
        conn.execute(
            """
            INSERT INTO projects (
                id, title, topic, genre, audience, tone, target_chapter_count,
                target_words_per_chapter, logline, synopsis, global_summary, status,
                privacy_mode, project_root_path, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft', ?, ?, ?, ?)
            """,
            (
                project_id,
                payload.title,
                payload.topic,
                payload.genre,
                payload.audience,
                payload.tone,
                payload.target_chapter_count,
                payload.target_words_per_chapter,
                payload.logline,
                payload.synopsis,
                payload.global_summary,
                int(payload.privacy_mode),
                str(root),
                now,
                now,
            ),
        )
        project = row_to_dict(conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone())
    return project


@app.get("/api/projects")
def list_projects() -> list[dict[str, Any]]:
    with connect() as conn:
        return rows_to_dicts(conn.execute("SELECT * FROM projects ORDER BY updated_at DESC").fetchall())


@app.get("/api/projects/{project_id}")
def get_project(project_id: str) -> dict[str, Any]:
    return require_project(project_id)


@app.patch("/api/projects/{project_id}")
def update_project(project_id: str, payload: ProjectIn) -> dict[str, Any]:
    require_project(project_id)
    now = utc_now()
    with connect() as conn:
        conn.execute(
            """
            UPDATE projects
            SET title = ?, topic = ?, genre = ?, audience = ?, tone = ?,
                target_chapter_count = ?, target_words_per_chapter = ?,
                logline = ?, synopsis = ?, global_summary = ?, privacy_mode = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (
                payload.title,
                payload.topic,
                payload.genre,
                payload.audience,
                payload.tone,
                payload.target_chapter_count,
                payload.target_words_per_chapter,
                payload.logline,
                payload.synopsis,
                payload.global_summary,
                int(payload.privacy_mode),
                now,
                project_id,
            ),
        )
        return row_to_dict(conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone())


@app.delete("/api/projects/{project_id}")
def delete_project(project_id: str, payload: DeleteProjectIn) -> dict[str, bool]:
    project = require_project(project_id)
    expected_password = os.getenv("AI_NOVEL_DELETE_PASSWORD") or project["title"]
    if payload.password != expected_password:
        raise HTTPException(status_code=403, detail="Delete password is incorrect")
    with connect() as conn:
        conn.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    return {"ok": True}


def require_chapter(project_id: str, chapter_id: str) -> dict[str, Any]:
    require_project(project_id)
    with connect() as conn:
        chapter = row_to_dict(
            conn.execute(
                "SELECT * FROM chapters WHERE id = ? AND project_id = ?",
                (chapter_id, project_id),
            ).fetchone()
        )
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found in project")
    return chapter


def create_structured_record(
    project_id: str,
    resource: str,
    title: str,
    category: str,
    content: str,
    payload: dict[str, Any],
    status: str = "active",
) -> dict[str, Any]:
    table = table_for_resource(resource)
    now = utc_now()
    record_id = new_id()
    with connect() as conn:
        conn.execute(
            f"""
            INSERT INTO {table} (id, project_id, title, category, content, payload, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (record_id, project_id, title, category, content, json.dumps(payload, ensure_ascii=False), status, now, now),
        )
        record = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE id = ?", (record_id,)).fetchone())
    sync_record_to_wiki(project_id, resource, record)
    return record


@app.post("/api/projects/{project_id}/chapters")
def create_chapter(project_id: str, payload: ChapterIn) -> dict[str, Any]:
    require_project(project_id)
    chapter_id = new_id()
    now = utc_now()
    word_count = len(payload.draft)
    with connect() as conn:
        conn.execute(
            """
            INSERT INTO chapters (
                id, project_id, outline_id, chapter_number, title, brief, draft,
                summary, word_count, status, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                chapter_id,
                project_id,
                payload.outline_id,
                payload.chapter_number,
                payload.title,
                payload.brief,
                payload.draft,
                payload.summary,
                word_count,
                payload.status,
                now,
                now,
            ),
        )
        chapter = row_to_dict(conn.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,)).fetchone())
    write_chapter_snapshot(project_id, chapter)
    return chapter


@app.get("/api/projects/{project_id}/chapters")
def list_chapters(project_id: str) -> list[dict[str, Any]]:
    require_project(project_id)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(
                "SELECT * FROM chapters WHERE project_id = ? ORDER BY chapter_number",
                (project_id,),
            ).fetchall()
        )


@app.get("/api/projects/{project_id}/chapters/{chapter_id}")
def get_chapter(project_id: str, chapter_id: str) -> dict[str, Any]:
    return require_chapter(project_id, chapter_id)


@app.patch("/api/projects/{project_id}/chapters/{chapter_id}")
def update_chapter(project_id: str, chapter_id: str, payload: ChapterIn) -> dict[str, Any]:
    require_chapter(project_id, chapter_id)
    now = utc_now()
    with connect() as conn:
        conn.execute(
            """
            UPDATE chapters
            SET outline_id = ?, chapter_number = ?, title = ?, brief = ?, draft = ?,
                summary = ?, word_count = ?, status = ?, updated_at = ?
            WHERE id = ? AND project_id = ?
            """,
            (
                payload.outline_id,
                payload.chapter_number,
                payload.title,
                payload.brief,
                payload.draft,
                payload.summary,
                len(payload.draft),
                payload.status,
                now,
                chapter_id,
                project_id,
            ),
        )
        chapter = row_to_dict(conn.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,)).fetchone())
    write_chapter_snapshot(project_id, chapter)
    return chapter


@app.delete("/api/projects/{project_id}/chapters/{chapter_id}")
def delete_chapter(project_id: str, chapter_id: str) -> dict[str, bool]:
    chapter = require_chapter(project_id, chapter_id)
    with connect() as conn:
        conn.execute("DELETE FROM chapter_versions WHERE project_id = ? AND chapter_id = ?", (project_id, chapter_id))
        conn.execute("DELETE FROM chapters WHERE id = ? AND project_id = ?", (chapter_id, project_id))
    try:
        snapshot = project_root(project_id) / "manuscript" / f"chapter-{int(chapter['chapter_number']):03}.md"
        if snapshot.exists():
            snapshot.unlink()
    except (OSError, TypeError, ValueError):
        pass
    return {"ok": True}


@app.post("/api/projects/{project_id}/chapters/{chapter_id}/versions")
def create_chapter_version(project_id: str, chapter_id: str, payload: VersionIn) -> dict[str, Any]:
    require_chapter(project_id, chapter_id)
    version_id = new_id()
    now = utc_now()
    with connect() as conn:
        conn.execute(
            """
            INSERT INTO chapter_versions (id, project_id, chapter_id, label, content, model, context_summary, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (version_id, project_id, chapter_id, payload.label, payload.content, payload.model, payload.context_summary, now),
        )
        return row_to_dict(conn.execute("SELECT * FROM chapter_versions WHERE id = ?", (version_id,)).fetchone())


@app.get("/api/projects/{project_id}/chapters/{chapter_id}/versions")
def list_chapter_versions(project_id: str, chapter_id: str) -> list[dict[str, Any]]:
    require_chapter(project_id, chapter_id)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(
                "SELECT * FROM chapter_versions WHERE project_id = ? AND chapter_id = ? ORDER BY created_at DESC",
                (project_id, chapter_id),
            ).fetchall()
        )


@app.post("/api/projects/{project_id}/chapters/{chapter_id}/versions/{version_id}/select")
def select_chapter_version(project_id: str, chapter_id: str, version_id: str) -> dict[str, Any]:
    require_chapter(project_id, chapter_id)
    with connect() as conn:
        version = row_to_dict(
            conn.execute(
                "SELECT * FROM chapter_versions WHERE id = ? AND project_id = ? AND chapter_id = ?",
                (version_id, project_id, chapter_id),
            ).fetchone()
        )
        if not version:
            raise HTTPException(status_code=404, detail="Version not found in chapter")
        conn.execute(
            """
            UPDATE chapters
            SET draft = ?, selected_version_id = ?, word_count = ?, updated_at = ?
            WHERE id = ? AND project_id = ?
            """,
            (version["content"], version_id, len(version["content"]), utc_now(), chapter_id, project_id),
        )
        chapter = row_to_dict(conn.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,)).fetchone())
    write_chapter_snapshot(project_id, chapter)
    return chapter


@app.post("/api/projects/{project_id}/chapters/{chapter_id}/finalize")
def finalize_chapter(project_id: str, chapter_id: str) -> dict[str, Any]:
    chapter = require_chapter(project_id, chapter_id)
    summary = chapter["summary"] or f"第 {chapter['chapter_number']} 章定稿：{chapter['draft'][:80]}"
    now = utc_now()
    with connect() as conn:
        conn.execute(
            "UPDATE chapters SET status = 'final', summary = ?, updated_at = ? WHERE id = ? AND project_id = ?",
            (summary, now, chapter_id, project_id),
        )
        memory_id = new_id()
        conn.execute(
            """
            INSERT INTO memory_items (id, project_id, title, category, content, payload, status, created_at, updated_at)
            VALUES (?, ?, ?, 'chapter_summary', ?, '{}', 'approved', ?, ?)
            """,
            (memory_id, project_id, f"第 {chapter['chapter_number']} 章摘要", summary, now, now),
        )
        updated = row_to_dict(conn.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,)).fetchone())
    append_wiki_page(project_id, "global-summary.md", f"\n\n## 第 {chapter['chapter_number']} 章\n\n{summary}", chapter_id)
    sync_chapter_memory_to_wiki(project_id, updated)
    return updated


def sync_chapter_memory_to_wiki(project_id: str, chapter: dict[str, Any]) -> None:
    chapter_number = chapter.get("chapter_number") or 0
    title = chapter.get("title") or f"第 {chapter_number} 章"
    summary = chapter.get("summary") or (chapter.get("draft") or "")[:120]
    draft = chapter.get("draft") or ""
    create_structured_record(
        project_id,
        "timeline-events",
        f"第 {chapter_number} 章事件",
        "chapter_event",
        summary,
        {
            "event_time": f"第 {chapter_number} 章",
            "chapter": title,
            "characters": "",
            "cause": chapter.get("brief") or "",
            "status": "已定稿",
        },
        "active",
    )
    if "伏笔" in draft or "埋下" in draft:
        create_structured_record(
            project_id,
            "foreshadowings",
            f"第 {chapter_number} 章伏笔",
            "open",
            "章节定稿时自动提取的伏笔线索。",
            {
                "setup_chapter": title,
                "payoff_chapter": "",
                "status": "未回收",
                "related_characters": "",
                "hint": "章节中出现伏笔或埋线提示。",
            },
            "open",
        )
    append_wiki_page(project_id, "chapters/index.md", f"\n\n## {title}\n\n{summary}", str(chapter.get("id") or ""))


def write_chapter_snapshot(project_id: str, chapter: dict[str, Any]) -> None:
    root = project_root(project_id)
    path = root / "manuscript" / f"chapter-{int(chapter['chapter_number']):03}.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(f"# {chapter['title'] or '未命名章节'}\n\n{chapter['draft'] or ''}", encoding="utf-8")


def upsert_wiki_page(project_id: str, relative_path: str, content: str, source_chapter_id: str = "") -> dict[str, Any]:
    target = safe_wiki_path(project_id, relative_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    now = utc_now()
    with connect() as conn:
        existing = row_to_dict(
            conn.execute(
                "SELECT * FROM wiki_pages WHERE project_id = ? AND path = ?",
                (project_id, relative_path),
            ).fetchone()
        )
        if existing:
            conn.execute(
                "UPDATE wiki_pages SET content = ?, updated_at = ? WHERE id = ?",
                (content, now, existing["id"]),
            )
            page_id = existing["id"]
        else:
            page_id = new_id()
            conn.execute(
                """
                INSERT INTO wiki_pages (id, project_id, path, title, content, updated_at)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (page_id, project_id, relative_path, Path(relative_path).stem, content, now),
            )
        conn.execute(
            """
            INSERT INTO wiki_page_revisions (id, project_id, path, content, source_chapter_id, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (new_id(), project_id, relative_path, content, source_chapter_id, now),
        )
        return row_to_dict(conn.execute("SELECT * FROM wiki_pages WHERE id = ?", (page_id,)).fetchone())


def append_wiki_page(project_id: str, relative_path: str, content: str, source_chapter_id: str = "") -> dict[str, Any]:
    target = safe_wiki_path(project_id, relative_path)
    previous = target.read_text(encoding="utf-8") if target.exists() else f"# {Path(relative_path).stem}\n"
    return upsert_wiki_page(project_id, relative_path, previous + content, source_chapter_id)


def delete_wiki_page(project_id: str, relative_path: str) -> None:
    target = safe_wiki_path(project_id, relative_path)
    if target.exists():
        target.unlink()
    with connect() as conn:
        conn.execute("DELETE FROM wiki_pages WHERE project_id = ? AND path = ?", (project_id, relative_path))


def safe_wiki_filename(value: str, fallback: str) -> str:
    cleaned = (value or fallback).strip().replace("\\", "-").replace("/", "-")
    return cleaned or fallback


def record_payload(record: dict[str, Any]) -> dict[str, Any]:
    payload = record.get("payload")
    return payload if isinstance(payload, dict) else {}


def markdown_for_record(resource: str, record: dict[str, Any]) -> str:
    payload = record_payload(record)
    title = record.get("title") or payload.get("name") or "未命名资料"
    lines = [f"# {title}", "", f"- 类型：{resource}", f"- 状态：{record.get('status') or 'active'}"]
    if record.get("category"):
        lines.append(f"- 分类：{record['category']}")
    for key, value in payload.items():
        if value not in ("", None, [], {}):
            lines.append(f"- {key}：{value}")
    if record.get("content"):
        lines.extend(["", "## 内容", "", str(record["content"])])
    return "\n".join(lines) + "\n"


def aggregate_markdown(project_id: str, resource: str, heading: str) -> str:
    table = table_for_resource(resource)
    with connect() as conn:
        records = rows_to_dicts(
            conn.execute(f"SELECT * FROM {table} WHERE project_id = ? ORDER BY created_at", (project_id,)).fetchall()
        )
    lines = [f"# {heading}", ""]
    for record in records:
        lines.extend([f"## {record.get('title') or '未命名'}", "", record.get("content") or markdown_for_record(resource, record), ""])
    return "\n".join(lines).strip() + "\n"


def outline_record_key(record: dict[str, Any]) -> str:
    payload = record_payload(record)
    category = str(record.get("category") or payload.get("scope") or "")
    if category == "global_outline" or payload.get("scope") == "global":
        return "global"
    chapter_id = str(payload.get("chapter_id") or "").strip()
    if chapter_id:
        return f"chapter_id:{chapter_id}"
    chapter_number = str(payload.get("chapter_number") or "").strip()
    if chapter_number:
        return f"chapter_number:{chapter_number}"
    chapter_title = str(payload.get("chapter_title") or record.get("title") or "").strip()
    return f"title:{chapter_title}" if chapter_title else f"record:{record.get('id')}"


def aggregate_outline_markdown(project_id: str) -> str:
    with connect() as conn:
        records = rows_to_dicts(
            conn.execute(
                "SELECT * FROM outlines WHERE project_id = ? ORDER BY created_at, updated_at",
                (project_id,),
            ).fetchall()
        )
    deduped: dict[str, dict[str, Any]] = {}
    for record in records:
        deduped[outline_record_key(record)] = record

    def sort_key(record: dict[str, Any]) -> tuple[int, int, str]:
        payload = record_payload(record)
        key = outline_record_key(record)
        if key == "global":
            return (0, 0, str(record.get("updated_at") or ""))
        try:
            number = int(str(payload.get("chapter_number") or "999999"))
        except ValueError:
            number = 999999
        return (1, number, str(record.get("updated_at") or ""))

    lines = ["# 总线大纲与章节大纲", ""]
    for record in sorted(deduped.values(), key=sort_key):
        payload = record_payload(record)
        title = record.get("title") or payload.get("chapter_title") or "未命名大纲"
        lines.extend([f"## {title}", "", record.get("content") or markdown_for_record("outlines", record), ""])
    return "\n".join(lines).strip() + "\n"


def sync_record_to_wiki(project_id: str, resource: str, record: dict[str, Any]) -> None:
    payload = record_payload(record)
    title = str(record.get("title") or payload.get("name") or record.get("id") or "未命名")
    filename = safe_wiki_filename(title, str(record.get("id") or "record"))
    if resource in {"character-profiles", "characters"}:
        upsert_wiki_page(project_id, f"characters/{filename}.md", markdown_for_record(resource, record))
        upsert_wiki_page(project_id, "characters/index.md", aggregate_markdown(project_id, "character-profiles", "角色档案"))
    elif resource == "character-relationships":
        upsert_wiki_page(project_id, "relationships.md", aggregate_markdown(project_id, "character-relationships", "角色关系图"))
    elif resource == "outlines":
        upsert_wiki_page(project_id, f"outlines/{filename}.md", markdown_for_record(resource, record))
        upsert_wiki_page(project_id, "outline.md", aggregate_outline_markdown(project_id))
    elif resource == "style-profiles":
        upsert_wiki_page(project_id, f"styles/{filename}.md", markdown_for_record(resource, record))
        upsert_wiki_page(project_id, "style.md", aggregate_markdown(project_id, "style-profiles", "写作风格档案"))
    elif resource == "timeline-events":
        upsert_wiki_page(project_id, "timeline.md", aggregate_markdown(project_id, "timeline-events", "时间线"))
    elif resource == "foreshadowings":
        upsert_wiki_page(project_id, "foreshadowing.md", aggregate_markdown(project_id, "foreshadowings", "伏笔管理"))
    elif resource == "taboo-rules":
        upsert_wiki_page(project_id, "taboo-rules.md", aggregate_markdown(project_id, "taboo-rules", "雷点规则"))
    elif resource == "knowledge-documents":
        upsert_wiki_page(project_id, f"knowledge/{filename}.md", markdown_for_record(resource, record))
        upsert_wiki_page(project_id, "knowledge/index.md", aggregate_markdown(project_id, "knowledge-documents", "知识库"))


def delete_record_from_wiki(project_id: str, resource: str, record: dict[str, Any]) -> None:
    payload = record_payload(record)
    title = str(record.get("title") or payload.get("name") or record.get("id") or "未命名")
    filename = safe_wiki_filename(title, str(record.get("id") or "record"))
    if resource == "outlines":
        delete_wiki_page(project_id, f"outlines/{filename}.md")
        upsert_wiki_page(project_id, "outline.md", aggregate_outline_markdown(project_id))


def list_records_for_context(project_id: str, resource: str, limit: int = 20) -> list[dict[str, Any]]:
    table = table_for_resource(resource)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(
                f"SELECT * FROM {table} WHERE project_id = ? ORDER BY updated_at DESC LIMIT ?",
                (project_id, limit),
            ).fetchall()
        )


def build_generation_context(project_id: str, chapter_id: str = "") -> dict[str, Any]:
    chapter = require_chapter(project_id, chapter_id) if chapter_id else None
    with connect() as conn:
        recent_chapters = rows_to_dicts(
            conn.execute(
                "SELECT id, chapter_number, title, brief, summary FROM chapters WHERE project_id = ? ORDER BY chapter_number DESC LIMIT 3",
                (project_id,),
            ).fetchall()
        )
        wiki_pages = rows_to_dicts(
            conn.execute(
                "SELECT path, title, content FROM wiki_pages WHERE project_id = ? ORDER BY updated_at DESC LIMIT 12",
                (project_id,),
            ).fetchall()
        )
    return {
        "chapter": chapter,
        "recent_chapters": recent_chapters,
        "characters": list_records_for_context(project_id, "character-profiles"),
        "relationships": list_records_for_context(project_id, "character-relationships"),
        "outlines": list_records_for_context(project_id, "outlines"),
        "styles": list_records_for_context(project_id, "style-profiles"),
        "timeline": list_records_for_context(project_id, "timeline-events"),
        "foreshadowings": list_records_for_context(project_id, "foreshadowings"),
        "taboo_rules": list_records_for_context(project_id, "taboo-rules"),
        "knowledge": list_records_for_context(project_id, "knowledge-documents"),
        "wiki_pages": wiki_pages,
    }


def trim_text(value: Any, max_chars: int = 1200) -> Any:
    if not isinstance(value, str):
        return value
    if len(value) <= max_chars:
        return value
    return value[:max_chars].rstrip() + f"\n...[已截断 {len(value) - max_chars} 字]"


def compact_value(value: Any, max_string_chars: int = 1000, max_items: int = 12, depth: int = 0) -> Any:
    if isinstance(value, str):
        return trim_text(value, max_string_chars)
    if isinstance(value, list):
        return [compact_value(item, max_string_chars, max_items, depth + 1) for item in value[:max_items]]
    if isinstance(value, dict):
        if depth >= 4:
            return trim_text(json.dumps(value, ensure_ascii=False), max_string_chars)
        return {
            key: compact_value(item, max_string_chars, max_items, depth + 1)
            for key, item in value.items()
            if key not in {"api_key", "generation_context"}
        }
    return value


def compact_record(record: dict[str, Any], content_limit: int = 700) -> dict[str, Any]:
    payload = record.get("payload") if isinstance(record.get("payload"), dict) else {}
    return {
        "id": record.get("id", ""),
        "title": record.get("title", ""),
        "category": record.get("category", ""),
        "status": record.get("status", ""),
        "content": trim_text(record.get("content", ""), content_limit),
        "payload": compact_value(payload, 420, 10),
    }


def compact_chapter_for_prompt(chapter: dict[str, Any] | None) -> dict[str, Any] | None:
    if not chapter:
        return None
    return {
        "id": chapter.get("id", ""),
        "chapter_number": chapter.get("chapter_number", 0),
        "title": chapter.get("title", ""),
        "brief": trim_text(chapter.get("brief", ""), 900),
        "summary": trim_text(chapter.get("summary", ""), 900),
        "status": chapter.get("status", ""),
        "draft_excerpt": trim_text(chapter.get("draft", ""), 2200),
    }


def compact_generation_context(context: dict[str, Any]) -> dict[str, Any]:
    return {
        "chapter": compact_chapter_for_prompt(context.get("chapter")),
        "recent_chapters": [
            {
                "id": chapter.get("id", ""),
                "chapter_number": chapter.get("chapter_number", 0),
                "title": chapter.get("title", ""),
                "brief": trim_text(chapter.get("brief", ""), 500),
                "summary": trim_text(chapter.get("summary", ""), 700),
            }
            for chapter in (context.get("recent_chapters") or [])[:3]
            if isinstance(chapter, dict)
        ],
        "characters": [compact_record(record, 600) for record in (context.get("characters") or [])[:12] if isinstance(record, dict)],
        "relationships": [compact_record(record, 500) for record in (context.get("relationships") or [])[:16] if isinstance(record, dict)],
        "outlines": [compact_record(record, 900) for record in (context.get("outlines") or [])[:12] if isinstance(record, dict)],
        "styles": [compact_record(record, 900) for record in (context.get("styles") or [])[:6] if isinstance(record, dict)],
        "timeline": [compact_record(record, 500) for record in (context.get("timeline") or [])[:16] if isinstance(record, dict)],
        "foreshadowings": [compact_record(record, 500) for record in (context.get("foreshadowings") or [])[:16] if isinstance(record, dict)],
        "taboo_rules": [compact_record(record, 500) for record in (context.get("taboo_rules") or [])[:16] if isinstance(record, dict)],
        "knowledge": [compact_record(record, 700) for record in (context.get("knowledge") or [])[:8] if isinstance(record, dict)],
        "wiki_pages": [
            {
                "path": page.get("path", ""),
                "title": page.get("title", ""),
                "content": trim_text(page.get("content", ""), 1000),
            }
            for page in (context.get("wiki_pages") or [])[:8]
            if isinstance(page, dict)
        ],
    }


def compact_payload_for_remote(workflow: str, payload: AiWorkflowIn) -> dict[str, Any]:
    data = payload.model_dump()
    data.pop("generation_context", None)
    if workflow in {"generate_chapter_draft", "revise_selection"}:
        data["current_draft"] = trim_text(data.get("current_draft", ""), 3500)
        data["selected_text"] = trim_text(data.get("selected_text", ""), 2500)
        data["prompt"] = trim_text(data.get("prompt", ""), 1600)
        if isinstance(data.get("style_profile"), dict):
            data["style_profile"] = compact_record(data["style_profile"], 800)
        if isinstance(data.get("style_profiles"), list):
            data["style_profiles"] = [
                {"id": item.get("id", ""), "title": item.get("title", "")}
                for item in data["style_profiles"][:8]
                if isinstance(item, dict)
            ]
    return compact_value(data, 1400, 12)


def max_tokens_for_config(config_payload: dict[str, Any]) -> int:
    try:
        value = int(config_payload.get("max_tokens") or 4000)
    except (TypeError, ValueError):
        value = 4000
    return max(128, min(value, 32000))


def request_timeout_for_workflow(workflow: str) -> int:
    generation_workflows = {
        "generate_chapter_draft",
        "generate_chapter_brief",
        "generate_outline",
        "revise_selection",
    }
    env_key = "AI_NOVEL_GENERATION_TIMEOUT_SECONDS" if workflow in generation_workflows else "AI_NOVEL_MODEL_TIMEOUT_SECONDS"
    default_timeout = 600 if workflow in generation_workflows else 90
    try:
        value = int(os.getenv(env_key) or default_timeout)
    except ValueError:
        value = default_timeout
    return max(30, min(value, 1800))


@app.post("/api/projects/{project_id}/wiki/write")
def wiki_write(project_id: str, payload: WikiWriteIn) -> dict[str, Any]:
    require_project(project_id)
    return upsert_wiki_page(project_id, payload.path, payload.content, payload.source_chapter_id)


@app.post("/api/projects/{project_id}/wiki/append")
def wiki_append(project_id: str, payload: WikiWriteIn) -> dict[str, Any]:
    require_project(project_id)
    return append_wiki_page(project_id, payload.path, payload.content, payload.source_chapter_id)


@app.get("/api/projects/{project_id}/wiki/read")
def wiki_read(project_id: str, path: str) -> dict[str, str]:
    require_project(project_id)
    target = safe_wiki_path(project_id, path)
    if not target.exists():
        raise HTTPException(status_code=404, detail="Wiki page not found")
    return {"path": path, "content": target.read_text(encoding="utf-8")}


@app.get("/api/projects/{project_id}/wiki/search")
def wiki_search(project_id: str, q: str = "") -> list[dict[str, Any]]:
    require_project(project_id)
    with connect() as conn:
        rows = conn.execute(
            """
            SELECT * FROM wiki_pages
            WHERE project_id = ? AND (path LIKE ? OR content LIKE ?)
            ORDER BY updated_at DESC
            """,
            (project_id, f"%{q}%", f"%{q}%"),
        ).fetchall()
        return rows_to_dicts(rows)


@app.get("/api/projects/{project_id}/wiki/revisions")
def wiki_revisions(project_id: str, path: str) -> list[dict[str, Any]]:
    require_project(project_id)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(
                "SELECT * FROM wiki_page_revisions WHERE project_id = ? AND path = ? ORDER BY created_at DESC",
                (project_id, path),
            ).fetchall()
        )


@app.get("/api/projects/{project_id}/wiki/lint")
def wiki_lint(project_id: str) -> dict[str, Any]:
    require_project(project_id)
    with connect() as conn:
        pages = rows_to_dicts(conn.execute("SELECT * FROM wiki_pages WHERE project_id = ?", (project_id,)).fetchall())
    orphan_pages = [page["path"] for page in pages if page["path"] != "index.md" and page["path"] not in " ".join(p["content"] for p in pages)]
    return {"orphan_pages": orphan_pages, "page_count": len(pages), "warnings": []}


@app.get("/api/projects/{project_id}/{resource}")
def list_generic(project_id: str, resource: str) -> list[dict[str, Any]]:
    table = table_for_resource(resource)
    require_project(project_id)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(f"SELECT * FROM {table} WHERE project_id = ? ORDER BY created_at DESC", (project_id,)).fetchall()
        )


@app.post("/api/projects/{project_id}/{resource}")
def create_generic(project_id: str, resource: str, payload: GenericIn) -> dict[str, Any]:
    table = table_for_resource(resource)
    require_project(project_id)
    record_id = new_id()
    now = utc_now()
    with connect() as conn:
        conn.execute(
            f"""
            INSERT INTO {table} (id, project_id, title, category, content, payload, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                record_id,
                project_id,
                payload.title,
                payload.category,
                payload.content,
                json.dumps(payload.payload, ensure_ascii=False),
                payload.status,
                now,
                now,
            ),
        )
        record = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE id = ?", (record_id,)).fetchone())
    sync_record_to_wiki(project_id, resource, record)
    return record


@app.patch("/api/projects/{project_id}/{resource}/{record_id}")
def update_generic(project_id: str, resource: str, record_id: str, payload: GenericIn) -> dict[str, Any]:
    table = table_for_resource(resource)
    require_project(project_id)
    now = utc_now()
    with connect() as conn:
        existing = conn.execute(f"SELECT * FROM {table} WHERE id = ? AND project_id = ?", (record_id, project_id)).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Record not found")
        conn.execute(
            f"""
            UPDATE {table}
            SET title = ?, category = ?, content = ?, payload = ?, status = ?, updated_at = ?
            WHERE id = ? AND project_id = ?
            """,
            (
                payload.title,
                payload.category,
                payload.content,
                json.dumps(payload.payload, ensure_ascii=False),
                payload.status,
                now,
                record_id,
                project_id,
            ),
        )
        record = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE id = ?", (record_id,)).fetchone())
    sync_record_to_wiki(project_id, resource, record)
    return record


@app.delete("/api/projects/{project_id}/{resource}/{record_id}")
def delete_generic(project_id: str, resource: str, record_id: str) -> dict[str, bool]:
    table = table_for_resource(resource)
    require_project(project_id)
    with connect() as conn:
        record = row_to_dict(
            conn.execute(f"SELECT * FROM {table} WHERE id = ? AND project_id = ?", (record_id, project_id)).fetchone()
        )
        if not record:
            raise HTTPException(status_code=404, detail="Record not found in project")
        conn.execute(f"DELETE FROM {table} WHERE id = ? AND project_id = ?", (record_id, project_id))
    delete_record_from_wiki(project_id, resource, record)
    return {"ok": True}


def table_for_resource(resource: str) -> str:
    table = GENERIC_TABLES.get(resource)
    if not table:
        raise HTTPException(status_code=404, detail="Unknown resource")
    return table


def local_proxy_timeout_detail(base_url: str, model: str) -> str:
    parsed = urllib.parse.urlparse(base_url)
    host = parsed.hostname or ""
    port = parsed.port
    target = f"{host}:{port}" if port else host
    if host in {"127.0.0.1", "localhost", "::1"}:
        return (
            f"远程模型连接失败：本地模型代理 {target} 已接收测试请求，但等待上游模型响应超时。"
            f"当前模型：{model}。请检查 cli-proxy-api、mihomo/系统代理、节点稳定性，以及该代理是否支持此模型名。"
        )
    return (
        f"远程模型连接失败：请求 {base_url} 超时。当前模型：{model}。"
        "请检查 Base URL、网络代理、服务商状态和模型名。"
    )


@app.post("/api/projects/{project_id}/ai/test-connection")
def test_model_connection(project_id: str, payload: ModelConnectionTestIn) -> dict[str, Any]:
    require_project(project_id)
    api_key = payload.api_key.strip()
    model = payload.model_name.strip()
    base_url = (payload.base_url or "https://api.openai.com/v1").rstrip("/")
    if not api_key or not model:
        raise HTTPException(status_code=400, detail="缺少 API Key 或 Model Name，无法测试远程模型连接。")

    request_body = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是模型连接测试程序。"},
            {"role": "user", "content": "请只回复 OK。"},
        ],
        "temperature": payload.temperature,
        "max_tokens": max(1, min(payload.max_tokens or 16, 64)),
    }
    request = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(request_body, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            data = json.loads(response.read().decode("utf-8"))
        _ = data["choices"][0]["message"]["content"]
        return {"ok": True, "model": model, "message": "远程模型连接成功。"}
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")[:400]
        raise HTTPException(status_code=502, detail=f"远程模型连接失败：HTTP {exc.code} {detail}") from exc
    except (TimeoutError, socket.timeout) as exc:
        raise HTTPException(status_code=502, detail=local_proxy_timeout_detail(base_url, model)) from exc
    except (urllib.error.URLError, KeyError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=502, detail=f"远程模型连接失败：{exc}") from exc


@app.post("/api/projects/{project_id}/ai/{workflow}")
def run_ai_workflow(project_id: str, workflow: str, payload: AiWorkflowIn) -> dict[str, Any]:
    require_project(project_id)
    context = build_generation_context(project_id, payload.chapter_id)
    output = run_model_or_stub(project_id, workflow, payload, context)
    output["context"] = context
    output_status = output.get("status") or "success"
    output_error = output.get("error") or ""
    now = utc_now()
    with connect() as conn:
        conn.execute(
            """
            INSERT INTO ai_runs (id, project_id, workflow, input_snapshot, output_text, model, status, error_message, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                new_id(),
                project_id,
                workflow,
                json.dumps(payload.model_dump(), ensure_ascii=False),
                output["text"],
                output["model"],
                output_status,
                output_error,
                now,
            ),
        )
    if workflow == "score_chapter" and payload.chapter_id:
        require_chapter(project_id, payload.chapter_id)
        with connect() as conn:
            conn.execute(
                """
                INSERT INTO chapter_scores (id, project_id, chapter_id, total_score, payload, created_at)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (new_id(), project_id, payload.chapter_id, output["score"], json.dumps(output, ensure_ascii=False), now),
            )
    if workflow == "generate_chapter_variants" and payload.chapter_id:
        for index in range(max(1, payload.count)):
            create_chapter_version(
                project_id,
                payload.chapter_id,
                VersionIn(label=f"AI 版本 {index + 1}", content=f"{output['text']}\n\n候选版本 {index + 1}。"),
            )
    return output


CHARACTER_STUBS: list[dict[str, str]] = [
    {
        "name": "沈照夜",
        "role": "前朝公主",
        "faction": "流亡旧臣",
        "appearance": "二十岁上下，常穿素色斗篷，眼神克制。",
        "traits": "冷静、警惕、重诺",
        "desire": "夺回被篡改的记忆与王朝真相",
        "fear": "再次失去重要之人的记忆",
        "mainline_relation": "围绕改写记忆古籍推进主线",
        "arc": "从被动流亡到主动承担改写历史的代价",
        "voice": "短句、克制、少用感叹",
        "related_chapters": "",
        "notes": "核心主角，可承接记忆古籍主线。",
    },
    {
        "name": "顾临舟",
        "role": "旧朝密探",
        "faction": "流亡情报网",
        "appearance": "三十岁左右，常戴灰色手套，左眼下有旧疤。",
        "traits": "谨慎、讽刺、守口如瓶",
        "desire": "找出出卖旧朝密档的人",
        "fear": "自己保护的人再次成为牺牲品",
        "mainline_relation": "掌握古籍上一次现世的线索，推动主角接近幕后势力。",
        "arc": "从只交换情报到愿意暴露身份保护同伴",
        "voice": "话少，常用反问和冷幽默试探对方。",
        "related_chapters": "",
        "notes": "用于补足情报线和行动线。",
    },
    {
        "name": "苏晚",
        "role": "禁书馆抄录员",
        "faction": "市立旧书馆",
        "appearance": "身形单薄，袖口常有墨渍，随身带铜边笔记本。",
        "traits": "敏锐、胆小但不懦弱、记忆力惊人",
        "desire": "证明禁书馆失踪案不是意外",
        "fear": "被人发现她能记住被改写前的片段",
        "mainline_relation": "能察觉记忆改写后的缝隙，帮助主角校验真相。",
        "arc": "从旁观记录者成长为主动保存真相的人",
        "voice": "语速快，细节多，紧张时会重复关键词。",
        "related_chapters": "",
        "notes": "适合承担知识库、记忆校验和悬疑线索。",
    },
    {
        "name": "谢无咎",
        "role": "现朝监察使",
        "faction": "监察司",
        "appearance": "黑衣金扣，站姿端正，目光像审讯灯。",
        "traits": "强势、克制、相信秩序",
        "desire": "阻止古籍造成更大范围的记忆污染",
        "fear": "秩序只是另一种被编造的谎言",
        "mainline_relation": "与主角立场相冲，却可能在关键章节成为临时同盟。",
        "arc": "从追捕者转为共同承担真相代价的见证者",
        "voice": "句子规整，不轻易承诺，一旦承诺必执行。",
        "related_chapters": "",
        "notes": "适合制造外部压力和价值观冲突。",
    },
]


def collect_existing_character_names(payload: AiWorkflowIn, context: dict[str, Any]) -> set[str]:
    names: set[str] = set()
    extra = getattr(payload, "model_extra", {}) or {}
    for name in extra.get("existing_character_names") or []:
        if isinstance(name, str) and name.strip():
            names.add(name.strip())
    for record in extra.get("existing_characters") or []:
        if not isinstance(record, dict):
            continue
        record_payload = record.get("payload") if isinstance(record.get("payload"), dict) else {}
        name = record_payload.get("name") or record.get("title")
        if isinstance(name, str) and name.strip():
            names.add(name.strip())
    for record in context.get("characters") or []:
        if not isinstance(record, dict):
            continue
        record_payload = record.get("payload") if isinstance(record.get("payload"), dict) else {}
        name = record_payload.get("name") or record.get("title")
        if isinstance(name, str) and name.strip():
            names.add(name.strip())
    return names


def character_stub_for_payload(payload: AiWorkflowIn, context: dict[str, Any]) -> dict[str, str]:
    existing_names = collect_existing_character_names(payload, context)
    character = next((item for item in CHARACTER_STUBS if item["name"] not in existing_names), CHARACTER_STUBS[-1])
    return {
        **character,
        "notes": f"{character['notes']} 根据提示生成：{payload.prompt or '新角色'}",
    }


def structured_output_for_workflow(workflow: str, payload: AiWorkflowIn, context: dict[str, Any]) -> Any:
    if workflow == "generate_characters":
        return character_stub_for_payload(payload, context)
    if workflow in {"generate_outline", "generate_chapter_brief"}:
        chapter = context.get("chapter") or {}
        return {
            "volume": "第一卷",
            "chapter_title": chapter.get("title") or "未命名章节",
            "chapter_goal": payload.prompt or "推进主角发现古籍代价",
            "main_conflict": "主角想获得真相，但每次改写都会损失记忆。",
            "key_events": "发现线索；遭遇阻碍；做出选择；留下新的悬念。",
            "emotional_rhythm": "压抑开场，中段紧张，结尾留钩子。",
            "foreshadowing": "古籍的代价尚未完全揭示。",
            "hook": "她发现自己忘记了一个本不该忘记的人。",
            "related_characters": "沈照夜",
            "completion_status": "草稿",
        }
    if workflow == "extract_timeline_events":
        return [
            {
                "event_time": "当前章节",
                "chapter": (context.get("chapter") or {}).get("title") or "",
                "characters": "主角",
                "cause": payload.content or payload.prompt,
                "status": "待确认",
            }
        ]
    if workflow == "check_taboo_rules":
        return {"risk_level": "低", "issues": [], "suggestion": "未发现明显雷点，可继续人工复核。"}
    return None


def build_stub_ai_output(
    workflow: str,
    payload: AiWorkflowIn,
    context: dict[str, Any] | None = None,
    error: str = "当前使用本地占位模型。",
) -> dict[str, Any]:
    context = context or {}
    titles = {
        "generate_setting": "小说设定",
        "generate_characters": "人物卡",
        "generate_outline": "总纲",
        "generate_chapter_directory": "章节目录",
        "generate_chapter_brief": "本章大纲",
        "generate_chapter_draft": "章节正文",
        "summarize_chapter": "章节摘要",
        "extract_memory": "记忆提取",
        "extract_timeline_events": "时间线提取",
        "extract_relationships": "关系变化",
        "check_consistency": "一致性检查",
        "check_taboo_rules": "雷点检查",
        "analyze_style_sample": "风格分析",
        "revise_selection": "改写结果",
        "score_chapter": "章节评分",
    }
    title = titles.get(workflow, workflow)
    structured = structured_output_for_workflow(workflow, payload, context)
    text = (
        json.dumps(structured, ensure_ascii=False, indent=2)
        if structured is not None
        else f"## {title}\n\n这是本地 MVP 的可编辑 AI 占位结果。输入提示：{payload.prompt or payload.content or '无'}"
    )
    score = 82 if workflow == "score_chapter" else 0
    return {
        "workflow": workflow,
        "model": "local-stub",
        "text": text,
        "score": score,
        "structured": structured,
        "status": "local",
        "error": error,
        "items": [{"title": title, "content": text}],
    }


def parse_structured_ai_text(text: str) -> Any:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:].strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None


def system_prompt_for_workflow(workflow: str) -> str:
    if workflow in {"generate_chapter_draft", "revise_selection"}:
        return (
            "你是专业的中文长篇小说创作助手。当前任务是生成或改写小说正文。"
            "只返回可直接放入章节编辑器的中文正文，不要返回 JSON、Markdown 标题、chapter_id、drafts 数组或多个版本。"
            "必须参考上下文中的 llmwiki 记忆、角色、大纲、时间线、伏笔、雷点和知识库。"
            "如果用户输入包含章节编号，以 chapter_number 为准，绝不要把 chapter_id 当成章节序号。"
        )
    if workflow in {"generate_outline", "generate_chapter_brief"}:
        return (
            "你是专业的中文长篇小说大纲编辑。当前任务是生成结构化章节大纲，只返回 JSON，不要包裹解释。"
            "必须围绕当前章节标题展开，章节目标、冲突、关键事件、伏笔和结尾钩子都要服务于该标题。"
            "返回字段 chapter_title 必须包含章节数，格式使用“第 N 章 · 章节名”。"
            "必须先检查输入与 llmwiki 上下文中的 outlines、timeline、wiki_pages、foreshadowings。"
            "不得生成与已有大纲、时间线或 llmwiki 页面相同或高度相似的事件；如果已有事件出现过，要设计新的推进、反转或后果。"
            "多章大纲必须让每章事件相互区分，不能用同一发现、追查、争执或反转重复填充。"
        )
    return "你是专业的中文长篇小说创作助手。需要结构化工作流时，只返回 JSON，不要包裹解释。"


def run_model_or_stub(project_id: str, workflow: str, payload: AiWorkflowIn, context: dict[str, Any]) -> dict[str, Any]:
    config = resolve_model_config(project_id, workflow)
    if not config:
        return build_stub_ai_output(
            workflow,
            payload,
            context,
            f"当前使用本地占位模型：未找到可用于 {workflow} 的远程模型配置。请在设置中保存模型并设为默认，或在任务路由中为该任务选择模型。",
        )

    config_payload = config.get("payload") if isinstance(config.get("payload"), dict) else {}
    api_key = str(config_payload.get("api_key") or "")
    base_url = str(config_payload.get("base_url") or "https://api.openai.com/v1").rstrip("/")
    model = str(config_payload.get("model_name") or config.get("title") or "")
    if not api_key or not model:
        return build_stub_ai_output(
            workflow,
            payload,
            context,
            f"模型配置“{config.get('title') or '未命名模型'}”缺少 API Key 或 Model Name。",
        )

    remote_payload = compact_payload_for_remote(workflow, payload)
    remote_context = compact_generation_context(context)
    request_body = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt_for_workflow(workflow)},
            {
                "role": "user",
                "content": (
                    f"工作流：{workflow}\n\n"
                    f"输入：{json.dumps(remote_payload, ensure_ascii=False)}\n\n"
                    f"llmwiki 与写作上下文：{json.dumps(remote_context, ensure_ascii=False)}"
                ),
            },
        ],
        "temperature": float(config_payload.get("temperature") or 0.7),
        "max_tokens": max_tokens_for_config(config_payload),
    }
    request = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(request_body, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    timeout_seconds = request_timeout_for_workflow(workflow)
    try:
        with urllib.request.urlopen(request, timeout=timeout_seconds) as response:
            data = json.loads(response.read().decode("utf-8"))
        text = data["choices"][0]["message"]["content"]
        return {
            "workflow": workflow,
            "model": model,
            "text": text,
            "score": 0,
            "structured": parse_structured_ai_text(text),
            "status": "success",
            "error": "",
            "items": [{"title": workflow, "content": text}],
        }
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")[:1000].strip()
        summary = f"HTTP {exc.code}: {detail or exc.reason}"
        fallback = build_stub_ai_output(workflow, payload, context, "远程模型调用失败。")
        fallback["status"] = "fallback"
        fallback["error"] = summary
        return fallback
    except TimeoutError as exc:
        fallback = build_stub_ai_output(workflow, payload, context, "远程模型调用超时。")
        fallback["status"] = "fallback"
        fallback["error"] = (
            f"远程模型仍可能在生成，但 {timeout_seconds} 秒内暂未返回结果。"
            "这通常不是提示词或网络配置错误；可稍后重试，或通过 AI_NOVEL_GENERATION_TIMEOUT_SECONDS 继续放宽等待时间。"
            f"原始错误：{exc}"
        )
        return fallback
    except (urllib.error.URLError, KeyError, json.JSONDecodeError, TimeoutError) as exc:
        fallback = build_stub_ai_output(workflow, payload, context, "远程模型调用失败。")
        fallback["status"] = "fallback"
        fallback["error"] = str(exc)
        return fallback


def resolve_model_config(project_id: str, workflow: str) -> dict[str, Any] | None:
    with connect() as conn:
        route = row_to_dict(
            conn.execute(
                """
                SELECT * FROM model_task_routes
                WHERE project_id = ? AND (category = ? OR title = ?)
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (project_id, workflow, workflow),
            ).fetchone()
        )
        if route and route.get("content"):
            config = row_to_dict(
                conn.execute(
                    "SELECT * FROM model_configs WHERE project_id = ? AND id = ?",
                    (project_id, route["content"]),
                ).fetchone()
            )
            if config:
                return config
        configs = rows_to_dicts(
            conn.execute(
                "SELECT * FROM model_configs WHERE project_id = ? ORDER BY created_at DESC",
                (project_id,),
            ).fetchall()
        )
        default_config = next(
            (
                config
                for config in configs
                if isinstance(config.get("payload"), dict) and config["payload"].get("is_default")
            ),
            None,
        )
        return default_config or (configs[0] if configs else None)


def project_chapters_for_export(project_id: str) -> list[dict[str, Any]]:
    require_project(project_id)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(
                "SELECT * FROM chapters WHERE project_id = ? ORDER BY chapter_number",
                (project_id,),
            ).fetchall()
        )


def render_markdown(project_id: str) -> str:
    project = require_project(project_id)
    chapters = project_chapters_for_export(project_id)
    lines = [f"# {project['title']}", ""]
    if project.get("synopsis"):
        lines.extend(["## 简介", "", project["synopsis"], ""])
    for chapter in chapters:
        lines.extend([f"## 第 {chapter['chapter_number']} 章 {chapter['title']}", "", chapter["draft"] or "", ""])
    return "\n".join(lines)


@app.get("/api/projects/{project_id}/export/markdown", response_class=PlainTextResponse)
def export_markdown(project_id: str) -> str:
    content = render_markdown(project_id)
    export_path = project_root(project_id) / "exports" / "novel.md"
    export_path.write_text(content, encoding="utf-8")
    return content


@app.get("/api/projects/{project_id}/export/txt", response_class=PlainTextResponse)
def export_txt(project_id: str) -> str:
    content = render_markdown(project_id).replace("#", "")
    export_path = project_root(project_id) / "exports" / "novel.txt"
    export_path.write_text(content, encoding="utf-8")
    return content


@app.get("/api/projects/{project_id}/export/docx")
def export_docx(project_id: str) -> Response:
    content = render_markdown(project_id)
    try:
        from docx import Document

        buffer = BytesIO()
        document = Document()
        for line in content.splitlines():
            if line.startswith("# "):
                document.add_heading(line.removeprefix("# "), level=1)
            elif line.startswith("## "):
                document.add_heading(line.removeprefix("## "), level=2)
            else:
                document.add_paragraph(line)
        document.save(buffer)
        data = buffer.getvalue()
    except Exception:
        data = content.encode("utf-8")
    (project_root(project_id) / "exports" / "novel.docx").write_bytes(data)
    return Response(data, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/api/projects/{project_id}/export/pdf")
def export_pdf(project_id: str) -> Response:
    content = render_markdown(project_id)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        _width, height = A4
        y = height - 48
        for line in content.splitlines():
            if y < 48:
                pdf.showPage()
                y = height - 48
            pdf.drawString(48, y, line[:92])
            y -= 18
        pdf.save()
        data = buffer.getvalue()
    except Exception:
        data = content.encode("utf-8")
    (project_root(project_id) / "exports" / "novel.pdf").write_bytes(data)
    return Response(data, media_type="application/pdf")


@app.get("/api/projects/{project_id}/export/epub")
def export_epub(project_id: str) -> Response:
    content = render_markdown(project_id)
    try:
        from ebooklib import epub

        project = require_project(project_id)
        book = epub.EpubBook()
        book.set_identifier(project_id)
        book.set_title(project["title"])
        book.set_language("zh-CN")
        chapter = epub.EpubHtml(title=project["title"], file_name="novel.xhtml", lang="zh-CN")
        chapter.content = "<pre>" + content.replace("&", "&amp;").replace("<", "&lt;") + "</pre>"
        book.add_item(chapter)
        book.toc = (chapter,)
        book.spine = ["nav", chapter]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        buffer = BytesIO()
        epub.write_epub(buffer, book)
        data = buffer.getvalue()
    except Exception:
        data = content.encode("utf-8")
    (project_root(project_id) / "exports" / "novel.epub").write_bytes(data)
    return Response(data, media_type="application/epub+zip")
