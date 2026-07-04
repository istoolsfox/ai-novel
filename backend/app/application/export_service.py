"""Application service for manuscript exports and deliverability manifests."""
from __future__ import annotations

import json
import re
from io import BytesIO
from typing import Any

from ..infrastructure.database import connect, rows_to_dicts
from ..infrastructure.storage import project_root, require_project
from ..workflows.generation import clean_chapter_title


EXPORT_FORMATS = ("markdown", "txt", "docx", "pdf", "epub")


def project_chapters_for_export(project_id: str) -> list[dict[str, Any]]:
    require_project(project_id)
    with connect() as conn:
        return rows_to_dicts(
            conn.execute(
                "SELECT * FROM chapters WHERE project_id = ? ORDER BY chapter_number",
                (project_id,),
            ).fetchall()
        )


def render_markdown(project_id: str) -> str:
    project = require_project(project_id)
    chapters = project_chapters_for_export(project_id)
    lines = [f"# {project['title']}", ""]
    if project.get("synopsis"):
        lines.extend(["## 简介", "", project["synopsis"], ""])
    for chapter in chapters:
        lines.extend([
            f"## 第 {chapter['chapter_number']} 章 {clean_chapter_title(chapter)}",
            "",
            strip_embedded_chapter_heading(chapter),
            "",
        ])
    return "\n".join(lines)


def strip_embedded_chapter_heading(chapter: dict[str, Any]) -> str:
    draft = (chapter.get("draft") or "").lstrip()
    if not draft:
        return ""
    title = clean_chapter_title(chapter)
    chapter_number = int(chapter.get("chapter_number") or 0)
    heading_pattern = rf"^(?:第\s*{chapter_number}\s*章(?:\s*[·:：-]\s*{re.escape(title)})?|{re.escape(title)})\s*\n+"
    return re.sub(heading_pattern, "", draft, count=1).lstrip()


def build_export_manifest(project_id: str) -> dict[str, Any]:
    project = require_project(project_id)
    chapters = project_chapters_for_export(project_id)
    total_words = sum(int(chapter.get("word_count") or 0) for chapter in chapters)
    final_chapters = [
        chapter
        for chapter in chapters
        if chapter.get("status") in {"final", "finalized"}
    ]
    target_count = int(project.get("target_chapter_count") or len(chapters) or 0)
    quality_scores = [
        float(chapter.get("quality_score") or 0)
        for chapter in chapters
        if float(chapter.get("quality_score") or 0) > 0
    ]
    average_quality = round(sum(quality_scores) / len(quality_scores), 1) if quality_scores else 0
    unfinished = [
        {"chapter_number": chapter.get("chapter_number"), "title": chapter.get("title") or ""}
        for chapter in chapters
        if chapter.get("status") not in {"final", "finalized"}
    ]
    low_quality = [
        {
            "chapter_number": chapter.get("chapter_number"),
            "title": chapter.get("title") or "",
            "quality_score": chapter.get("quality_score") or 0,
        }
        for chapter in chapters
        if chapter.get("quality_score") and float(chapter.get("quality_score") or 0) < 70
    ]
    missing_numbers = sorted(
        set(range(1, target_count + 1))
        - {int(chapter.get("chapter_number") or 0) for chapter in chapters}
    )
    deliverable = bool(
        target_count
        and len(chapters) >= target_count
        and len(final_chapters) >= target_count
        and not unfinished
        and not low_quality
        and not missing_numbers
    )
    manifest = {
        "project_id": project_id,
        "title": project["title"],
        "target_chapter_count": target_count,
        "chapter_count": len(chapters),
        "final_chapter_count": len(final_chapters),
        "total_words": total_words,
        "average_quality_score": average_quality,
        "deliverable": deliverable,
        "missing_chapter_numbers": missing_numbers,
        "unfinished_chapters": unfinished,
        "low_quality_chapters": low_quality,
        "exports": {
            "markdown": "exports/novel.md",
            "txt": "exports/novel.txt",
            "docx": "exports/novel.docx",
            "pdf": "exports/novel.pdf",
            "epub": "exports/novel.epub",
        },
    }
    manifest_path = project_root(project_id) / "exports" / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def export_markdown_file(project_id: str) -> str:
    content = render_markdown(project_id)
    (project_root(project_id) / "exports" / "novel.md").write_text(content, encoding="utf-8")
    return content


def export_txt_file(project_id: str) -> str:
    content = render_markdown(project_id).replace("#", "")
    (project_root(project_id) / "exports" / "novel.txt").write_text(content, encoding="utf-8")
    return content


def export_docx_file(project_id: str) -> bytes:
    content = render_markdown(project_id)
    try:
        from docx import Document

        buffer = BytesIO()
        document = Document()
        for line in content.splitlines():
            if line.startswith("# "):
                document.add_heading(line.removeprefix("# "), level=1)
            elif line.startswith("## "):
                document.add_heading(line.removeprefix("## "), level=2)
            else:
                document.add_paragraph(line)
        document.save(buffer)
        data = buffer.getvalue()
    except Exception:
        data = content.encode("utf-8")
    (project_root(project_id) / "exports" / "novel.docx").write_bytes(data)
    return data


def export_pdf_file(project_id: str) -> bytes:
    content = render_markdown(project_id)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        _width, height = A4
        y = height - 48
        for line in content.splitlines():
            if y < 48:
                pdf.showPage()
                y = height - 48
            pdf.drawString(48, y, line[:92])
            y -= 18
        pdf.save()
        data = buffer.getvalue()
    except Exception:
        data = content.encode("utf-8")
    (project_root(project_id) / "exports" / "novel.pdf").write_bytes(data)
    return data


def export_epub_file(project_id: str) -> bytes:
    content = render_markdown(project_id)
    try:
        from ebooklib import epub

        project = require_project(project_id)
        book = epub.EpubBook()
        book.set_identifier(project_id)
        book.set_title(project["title"])
        book.set_language("zh-CN")
        chapter = epub.EpubHtml(title=project["title"], file_name="novel.xhtml", lang="zh-CN")
        chapter.content = "<pre>" + content.replace("&", "&amp;").replace("<", "&lt;") + "</pre>"
        book.add_item(chapter)
        book.toc = (chapter,)
        book.spine = ["nav", chapter]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        buffer = BytesIO()
        epub.write_epub(buffer, book)
        data = buffer.getvalue()
    except Exception:
        data = content.encode("utf-8")
    (project_root(project_id) / "exports" / "novel.epub").write_bytes(data)
    return data


def export_all_files(project_id: str) -> dict[str, Any]:
    export_markdown_file(project_id)
    export_txt_file(project_id)
    export_docx_file(project_id)
    export_pdf_file(project_id)
    export_epub_file(project_id)
    return build_export_manifest(project_id)
